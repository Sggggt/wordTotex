from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, List, Optional, Tuple, Union

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


BLOCK = Union[Paragraph, Table]


@dataclass
class ConverterConfig:
    include_preamble: bool = True
    table_border: bool = True


@dataclass
class ConversionResult:
    latex: str
    image_paths: List[Path]


class DocxToLatexConverter:
    def __init__(self, config: Optional[ConverterConfig] = None) -> None:
        self.config = config or ConverterConfig()
        self._image_dir: Optional[Path] = None
        self._saved_images: List[Path] = []
        self._tex_output_dir: Optional[Path] = None
        self._image_counter: int = 1

    def convert(
        self, docx_path: Union[str, Path], output_dir: Optional[Path] = None
    ) -> ConversionResult:
        docx_path = Path(docx_path)
        document = Document(docx_path)
        self._tex_output_dir = Path(output_dir) if output_dir else docx_path.parent
        self._image_dir = self._tex_output_dir / f"{docx_path.stem}_images"
        self._saved_images = []
        self._image_counter = 1

        lines: List[str] = []
        current_list: Optional[str] = None

        for block in self._iter_block_items(document):
            list_type = self._get_list_type(block) if isinstance(block, Paragraph) else None
            if current_list and list_type != current_list:
                lines.append(f"\\end{{{current_list}}}")
                current_list = None

            if isinstance(block, Paragraph):
                paragraph_tex, opened_list = self._convert_paragraph(block, list_type)
                if opened_list and not current_list:
                    lines.append(f"\\begin{{{opened_list}}}")
                    current_list = opened_list
                if paragraph_tex:
                    lines.append(paragraph_tex)
            elif isinstance(block, Table):
                if current_list:
                    lines.append(f"\\end{{{current_list}}}")
                    current_list = None
                lines.extend(self._convert_table(block))

        if current_list:
            lines.append(f"\\end{{{current_list}}}")

        body = "\n".join(lines).strip() + "\n"
        latex = self._wrap_document(body) if self.config.include_preamble else body
        return ConversionResult(latex=latex, image_paths=self._saved_images.copy())

    def _convert_paragraph(
        self, paragraph: Paragraph, list_type: Optional[str]
    ) -> Tuple[str, Optional[str]]:
        text = self._build_runs(paragraph)
        if not text:
            return "", list_type

        style = paragraph.style.name if paragraph.style is not None else ""
        heading = self._heading_command(style)
        if heading:
            return f"{heading}{{{text}}}\n", None

        if list_type:
            return f"\\item {text}", list_type

        aligned = self._apply_alignment(text, paragraph.alignment)
        return f"{aligned}\n", None

    def _convert_table(self, table: Table) -> List[str]:
        rows_tex: List[str] = []
        num_cols = max(len(row.cells) for row in table.rows)
        alignment = "|".join(self._column_alignment(table, idx) for idx in range(num_cols))
        border = "|" if self.config.table_border else ""
        rows_tex.append(f"\\begin{{tabular}}{{{border}{alignment}{border}}}")
        rows_tex.append("\\hline")

        for row in table.rows:
            cells = [self._convert_cell(cell) for cell in row.cells]
            rows_tex.append(" & ".join(cells) + " \\\\")
            rows_tex.append("\\hline")

        rows_tex.append("\\end{tabular}\n")
        return rows_tex

    def _build_runs(self, paragraph: Paragraph) -> str:
        parts: List[str] = []
        for run in paragraph.runs:
            image_snippets = self._extract_images(run)
            parts.extend(image_snippets)
            content = self._escape_tex(run.text)
            if not content:
                continue

            parts.append(self._apply_run_formatting(content, run))
        return "".join(parts)

    def _escape_tex(self, text: str) -> str:
        if not text:
            return ""

        replacements = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        escaped = []
        for char in text.replace("\n", r"\\ "):
            escaped.append(replacements.get(char, char))
        return "".join(escaped)

    def _heading_command(self, style_name: str) -> Optional[str]:
        normalized = style_name.lower()
        if normalized.startswith("heading 1"):
            return "\\section"
        if normalized.startswith("heading 2"):
            return "\\subsection"
        if normalized.startswith("heading 3"):
            return "\\subsubsection"
        if normalized.startswith("heading 4"):
            return "\\paragraph"
        if normalized.startswith("heading 5"):
            return "\\subparagraph"
        return None

    def _get_list_type(self, paragraph: Paragraph) -> Optional[str]:
        style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
        if style_name.startswith("list bullet") or style_name.startswith("blockquote"):
            return "itemize"
        if style_name.startswith("list number"):
            return "enumerate"
        return None

    def _apply_alignment(self, text: str, alignment: Optional[int]) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return f"\\begin{{center}}{text}\\end{{center}}"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return f"\\begin{{flushright}}{text}\\end{{flushright}}"
        return text

    def _apply_run_formatting(self, content: str, run: Run) -> str:
        formatted = content

        font = run.font
        hex_color = None
        if font and font.color and font.color.rgb:
            rgb = font.color.rgb
            hex_color = getattr(rgb, "hex", None) or str(rgb).replace("0x", "").replace("#", "")
            if len(hex_color) == 8:  # strip alpha if present
                hex_color = hex_color[2:]
        if hex_color:
            formatted = f"\\textcolor[HTML]{{{hex_color}}}{{{formatted}}}"

        if run.underline:
            formatted = f"\\underline{{{formatted}}}"
        if run.italic:
            formatted = f"\\textit{{{formatted}}}"
        if run.bold:
            formatted = f"\\textbf{{{formatted}}}"

        if font and font.name and self._is_monospace(font.name):
            formatted = f"\\texttt{{{formatted}}}"

        size_cmd = self._size_command(font.size.pt if font and font.size else None)
        if size_cmd:
            formatted = f"{{{size_cmd} {formatted}}}"
        return formatted

    def _column_alignment(self, table: Table, col_idx: int) -> str:
        alignments: List[str] = []
        for row in table.rows:
            if col_idx >= len(row.cells):
                continue
            cell = row.cells[col_idx]
            alignment = None
            if cell.paragraphs:
                alignment = cell.paragraphs[0].alignment
            alignments.append(self._map_alignment(alignment))

        # pick the most common alignment in this column, fallback to left
        left = "l"
        if not alignments:
            return left
        counts = {"l": 0, "c": 0, "r": 0}
        for align in alignments:
            counts[align] = counts.get(align, 0) + 1
        return max(counts.items(), key=lambda item: item[1])[0]

    def _convert_cell(self, cell: _Cell) -> str:
        if not cell.paragraphs:
            return ""
        parts: List[str] = []
        for paragraph in cell.paragraphs:
            text = self._build_runs(paragraph)
            if not text:
                continue
            align_prefix = self._cell_alignment_prefix(paragraph.alignment)
            if align_prefix:
                text = f"{{{align_prefix} {text}}}"
            parts.append(text)
        return r" \\ ".join(parts) if parts else ""

    def _map_alignment(self, alignment: Optional[int]) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "c"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "r"
        return "l"

    def _cell_alignment_prefix(self, alignment: Optional[int]) -> str:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "\\centering"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "\\raggedleft"
        return ""

    def _is_monospace(self, font_name: str) -> bool:
        lowered = font_name.lower()
        return any(keyword in lowered for keyword in ("mono", "consolas", "courier", "code"))

    def _size_command(self, size_pt: Optional[float]) -> Optional[str]:
        if size_pt is None:
            return None
        if size_pt >= 18:
            return "\\Large"
        if size_pt >= 16:
            return "\\large"
        if size_pt <= 8:
            return "\\footnotesize"
        if size_pt <= 9:
            return "\\small"
        return None

    def _wrap_document(self, body: str) -> str:
        preamble = [
            r"\documentclass{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage{array}",
            r"\usepackage{xcolor}",
            r"\usepackage{hyperref}",
            r"\usepackage{graphicx}",
            r"\begin{document}",
        ]
        return "\n".join(preamble) + "\n\n" + body + "\n\\end{document}\n"

    def _iter_block_items(self, parent: Union[DocumentType, _Cell]) -> Iterator[BLOCK]:
        parent_elm = parent.element.body if isinstance(parent, DocumentType) else parent._tc
        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    def _extract_images(self, run: Run) -> List[str]:
        """Save images embedded in a run and return LaTeX includegraphics commands."""
        if self._tex_output_dir is None:
            return []

        snippets: List[str] = []
        pictures = run.element.xpath(".//pic:pic")
        for pic in pictures:
            blips = pic.xpath(".//a:blip")
            if not blips:
                continue
            embed = blips[0].get(qn("r:embed"))
            if not embed:
                continue

            image_part = run.part.related_parts.get(embed)
            if not image_part:
                continue

            ext = Path(str(image_part.partname)).suffix or ".png"
            image_name = f"image_{self._image_counter}{ext}"
            self._image_counter += 1
            image_dir = self._ensure_image_dir()
            image_path = image_dir / image_name
            image_path.write_bytes(image_part.blob)
            self._saved_images.append(image_path)

            rel_path = self._relative_image_path(image_path)
            snippets.append(f"\n\\includegraphics[width=\\linewidth]{{{rel_path}}}\n")

        return snippets

    def _ensure_image_dir(self) -> Path:
        if self._image_dir is None:
            raise ValueError("Image directory is not configured")
        self._image_dir.mkdir(parents=True, exist_ok=True)
        return self._image_dir

    def _relative_image_path(self, image_path: Path) -> str:
        base_dir = self._tex_output_dir or image_path.parent
        try:
            rel_path = image_path.relative_to(base_dir)
        except ValueError:
            rel_path = image_path
        return rel_path.as_posix()
